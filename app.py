# -*- coding: utf-8 -*-
"""
新员工培训 AI 助手 - 后端服务
功能：
  1. 从 kb/text 加载内部培训资料，做 BM25 检索
  2. 对时效性/外部信息自动联网核验（SerpAPI 优先，回退免费搜索）
  3. 调用 DeepSeek 流式生成带出处的回答
  4. 账号体系（手机号+密码注册/登录/登出/找回密码）
  5. 使用记录埋点（员工每次提问入库，供总部后台查看）
  6. 总部后台 API（用户管理、群公告、助手改名、DeepSeek 余额）
员工通过手机浏览器访问前端即可提问；总部主账号访问 /admin 进行管理。
"""
import os, re, json, time
from functools import wraps
from datetime import datetime, timedelta

from flask import (Flask, request, Response, session, jsonify,
                   redirect, render_template, send_from_directory)
from werkzeug.security import generate_password_hash, check_password_hash
import requests
from bs4 import BeautifulSoup

from db import db, make_db_uri, User, UsageLog, Setting, ResetRequest, Feedback, LearningProgress, QuizAttempt
from sqlalchemy import text

try:
    import jieba
    jieba.setLogLevel(20)
    def tokenize(text):
        return [t for t in jieba.lcut(text) if t.strip()]
except Exception:
    def tokenize(text):
        return list(text)

from rank_bm25 import BM25Okapi

# 本地开发时从 .env 读取密钥（手写解析，不依赖第三方包，避免安装/网络问题）
# 用 setdefault，不会覆盖 Render 等平台已注入的环境变量
def _load_local_env():
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
    if not os.path.exists(env_path):
        return
    with open(env_path, encoding="utf-8") as fh:
        for line in fh:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            k, v = k.strip(), v.strip().strip('"').strip("'")
            os.environ.setdefault(k, v)
_load_local_env()

BASE = os.path.dirname(os.path.abspath(__file__))
KB_DIR = os.path.join(BASE, "kb", "text")
ONB_DIR = os.path.join(BASE, "kb", "onboarding")
QUIZ_BANK_FILE = os.path.join(BASE, "kb", "quiz", "quiz-bank.json")
STATIC_DIR = os.path.join(BASE, "static")

# ---- 配置（可通过环境变量覆盖）----
SECRET_KEY = os.environ.get("SECRET_KEY", "dev-secret-change-in-prod")
ADMIN_PHONE = os.environ.get("ADMIN_PHONE", "")
ADMIN_DEFAULT_PASSWORD = os.environ.get("ADMIN_DEFAULT_PASSWORD", "")

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")
WEB_SEARCH_ENABLED = os.environ.get("WEB_SEARCH_ENABLED", "1") != "0"
SEARCH_API_KEY = os.environ.get("SEARCH_API_KEY", "")
SEARCH_ENGINE = os.environ.get("SEARCH_ENGINE", "baidu")
TOP_K = int(os.environ.get("TOP_K", "5"))
CHUNK_SIZE = int(os.environ.get("CHUNK_SIZE", "400"))
CHUNK_OVERLAP = int(os.environ.get("CHUNK_OVERLAP", "80"))

# ---- 知识库加载 ----
_chunks = []
_bm25 = None

def load_kb():
    global _chunks, _bm25
    _chunks = []
    files = [f for f in os.listdir(KB_DIR) if f.endswith(".txt")]
    for f in files:
        source = os.path.splitext(f)[0]
        path = os.path.join(KB_DIR, f)
        with open(path, encoding="utf-8", errors="ignore") as fh:
            text = fh.read()
        paras = [p.strip() for p in re.split(r"\n+", text) if p.strip()]
        buf = ""
        idx = 0
        for p in paras:
            if len(buf) + len(p) > CHUNK_SIZE and buf:
                _chunks.append({"id": f"{source}#{idx}", "source": source,
                                "text": buf, "tokens": tokenize(buf)})
                idx += 1
                buf = p
            else:
                buf = (buf + "\n" + p).strip()
        if buf:
            _chunks.append({"id": f"{source}#{idx}", "source": source,
                            "text": buf, "tokens": tokenize(buf)})
    corpus = [c["tokens"] for c in _chunks]
    _bm25 = BM25Okapi(corpus)
    print(f"[KB] 已加载 {len(files)} 篇文档, {len(_chunks)} 个片段")

def retrieve(query, k=TOP_K, max_per_source=2):
    if _bm25 is None:
        load_kb()
        if _bm25 is None:
            return []
    scores = _bm25.get_scores(tokenize(query))
    ranked = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)
    out = []
    per_source = {}
    for i in ranked:
        if scores[i] <= 0:
            continue
        c = _chunks[i]
        cnt = per_source.get(c["source"], 0)
        if cnt >= max_per_source:
            continue
        per_source[c["source"]] = cnt + 1
        out.append({"source": c["source"], "text": c["text"], "score": float(scores[i])})
        if len(out) >= k:
            break
    return out

# ---- 新员工必读板块 + 考试题库加载 ----
_onb_modules = []
_quiz_bank = []

# 必读板块与题库支持后台在线编辑：内容存 Setting 表（DB 持久化，部署不丢），
# 仓库内 JSON 文件仅作首次种子，一旦后台保存过则以 DB 为准。
ONB_DB_KEY = "onboarding_modules"
QUIZ_DB_KEY = "quiz_bank"
_onb_doc = {}    # 完整文档（含 meta）
_quiz_doc = {}
_onb_source = "file"
_quiz_source = "file"

def _read_setting(key):
    """读 Setting（自动兼容无 app context 的启动场景）"""
    try:
        with app.app_context():
            row = db.session.get(Setting, key)
            return row.value if row else None
    except Exception:
        return None

def _save_setting(key, value):
    with app.app_context():
        row = db.session.get(Setting, key)
        if row:
            row.value = value
        else:
            db.session.add(Setting(key=key, value=value))
        db.session.commit()

def load_onboarding():
    """加载「新员工必读」板块：DB（后台编辑版）优先，仓库文件兜底。"""
    global _onb_modules, _onb_doc, _onb_source
    data = None
    src = "db"
    raw = _read_setting(ONB_DB_KEY)
    if raw:
        try:
            data = json.loads(raw)
        except Exception:
            data = None
    if data is None:
        src = "file"
        try:
            path = os.path.join(ONB_DIR, "onboarding-modules.json")
            if os.path.exists(path):
                with open(path, encoding="utf-8") as fh:
                    data = json.load(fh)
        except Exception as e:
            print(f"[ONB] 文件加载失败: {e}")
    _onb_doc = data or {}
    _onb_modules = _onb_doc.get("modules", []) if isinstance(_onb_doc, dict) else []
    _onb_source = src
    print(f"[ONB] 已加载 {len(_onb_modules)} 个必读板块（来源：{src}）")

def load_quiz():
    """加载考试题库：DB（后台编辑版）优先，仓库文件兜底。"""
    global _quiz_bank, _quiz_doc, _quiz_source
    data = None
    src = "db"
    raw = _read_setting(QUIZ_DB_KEY)
    if raw:
        try:
            data = json.loads(raw)
        except Exception:
            data = None
    if data is None:
        src = "file"
        try:
            if os.path.exists(QUIZ_BANK_FILE):
                with open(QUIZ_BANK_FILE, encoding="utf-8") as fh:
                    data = json.load(fh)
        except Exception as e:
            print(f"[QUIZ] 文件加载失败: {e}")
    _quiz_doc = data or {}
    _quiz_bank = _quiz_doc.get("questions", []) if isinstance(_quiz_doc, dict) else []
    _quiz_source = src
    print(f"[QUIZ] 已加载 {len(_quiz_bank)} 道题（来源：{src}）")

# 各主题抽题配额（合计 20 题），确保多角度均衡
QUIZ_QUOTA = {"产品知识": 5, "抖音来客平台": 3, "门店准入": 3,
               "入驻流程": 3, "品牌授权": 2, "入驻审核": 4}
QUIZ_DURATION = 20  # 分钟
QUIZ_PER_SCORE = 5  # 每题分

def draw_questions(n=20):
    """从题库按主题均衡抽取 n 道题，返回题 id 列表（随机打乱顺序）。"""
    import random
    by_topic = {}
    for q in _quiz_bank:
        by_topic.setdefault(q.get("topic", ""), []).append(q)
    for t in by_topic:
        random.shuffle(by_topic[t])
    chosen = []
    for t, qn in QUIZ_QUOTA.items():
        pool = by_topic.get(t, [])
        chosen.extend([q["id"] for q in pool[:qn]])
    # 题库扩充后若配额不足，用剩余题补足
    if len(chosen) < n:
        picked = set(chosen)
        remain = [q["id"] for t in by_topic for q in by_topic[t] if q["id"] not in picked]
        random.shuffle(remain)
        chosen.extend(remain[: n - len(chosen)])
    random.shuffle(chosen)
    return chosen[:n]

def quiz_question_by_id(qid):
    for q in _quiz_bank:
        if q["id"] == qid:
            return q
    return None

def quiz_to_front(q, with_answer=False):
    """转换为前端结构；with_answer=True 时附带正确答案（考后显示）。"""
    opts = q.get("options", {})
    options = [{"key": k, "text": v} for k, v in opts.items()]
    out = {"id": q["id"], "topic": q.get("topic", ""), "stem": q["stem"], "options": options}
    if with_answer:
        out["answer"] = q.get("answer", "")
    return out

# ---- 时效性判断 ----
TIME_WORDS = ["最新", "今天", "今日", "今年", "去年", "2024", "2025", "2026", "政策", "规定",
              "法规", "法律", "补贴", "个税", "社保", "公积金", "价格", "报价", "活动", "截止",
              "新闻", "上市", "财报", "利率", "标准", "通知", "公告", "趋势", "行情", "工资",
              "最低工资", "放假", "假期", "节日", "限时", "新规", "调整"]

def is_time_sensitive(q):
    return any(w in q for w in TIME_WORDS)

# ---- 联网搜索 ----
def _search_serpapi(q, top=5):
    try:
        r = requests.get("https://serpapi.com/search.json",
                         params={"engine": SEARCH_ENGINE, "q": q, "api_key": SEARCH_API_KEY,
                                 "hl": "zh-cn", "gl": "cn"}, timeout=15)
        data = r.json()
        out = []
        for item in data.get("organic_results", [])[:top]:
            out.append({"title": item.get("title", ""),
                        "url": item.get("link") or item.get("url", ""),
                        "snippet": item.get("snippet", "")})
        return out
    except Exception as e:
        print(f"[WEB] serpapi 失败: {e}")
        return []

def _scrape_baidu(q, top=5):
    try:
        h = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120 Safari/537.36",
             "Accept-Language": "zh-CN,zh;q=0.9"}
        r = requests.get("https://www.baidu.com/s", params={"wd": q, "rn": str(top)}, headers=h, timeout=12)
        soup = BeautifulSoup(r.text, "html.parser")
        out = []
        for c in soup.select("div.c-container"):
            a = c.select_one("h3 a") or c.select_one("a")
            if not a:
                continue
            title = a.get_text(strip=True)
            link = a.get("href", "")
            sn = c.select_one("div.c-abstract, span.content-right_8Zs40")
            snippet = sn.get_text(strip=True) if sn else ""
            if title:
                out.append({"title": title, "url": link, "snippet": snippet})
            if len(out) >= top:
                break
        return out
    except Exception as e:
        print(f"[WEB] baidu 失败: {e}")
        return []

def _scrape_sogou(q, top=5):
    try:
        h = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120 Safari/537.36"}
        r = requests.get("https://www.sogou.com/web", params={"query": q}, headers=h, timeout=12)
        soup = BeautifulSoup(r.text, "html.parser")
        out = []
        for c in soup.select("div.vrwrap, div.rb"):
            a = c.select_one("h3 a")
            if not a:
                continue
            title = a.get_text(strip=True)
            link = a.get("href", "")
            sn = c.select_one("div.text-layout, p.str_info, div.fz-mid")
            snippet = sn.get_text(strip=True) if sn else ""
            if title:
                out.append({"title": title, "url": link, "snippet": snippet})
            if len(out) >= top:
                break
        return out
    except Exception as e:
        print(f"[WEB] sogou 失败: {e}")
        return []

def web_search(q, top=5):
    if not WEB_SEARCH_ENABLED:
        return []
    if SEARCH_API_KEY:
        res = _search_serpapi(q, top)
        if res:
            return res
        print("[WEB] serpapi 无结果，回退免费搜索")
    for fn in (_scrape_baidu, _scrape_sogou):
        try:
            res = fn(q, top)
            if res:
                return res
        except Exception as e:
            print(f"[WEB] 引擎异常: {e}")
    return []

# ---- 构建提示（助手名动态化，可在后台改名）----
SYSTEM_PROMPT_TEMPLATE = """你是「{name}」，名字叫小黄。你的本职工作是帮助公司新人解答入职后的各类业务与流程问题。

# 身份与风格
- 姓名：小黄
- 职业背景：八年产品经理，有丰富的公司管理经验，对市场风向与风口具有敏锐的感受。
- 思维习惯：面对用户问题时，能发散不同方向的思维；除了解答问题本身，会主动给出下一步动作的专业建议与专业看法（比如「你可以先…再…」「这类情况建议同步报备…」「下一步可以考虑…」）。
- 对话风格：温柔、理性、专业。不浮夸、不说教、有同理心；用词平实但有分量。

# 回答规则
1. 优先依据【内部资料库】内容回答公司相关信息（品牌、产品、卖点、话术、办证流程、拓店标准等）。
2. 当问题涉及时效性/外部信息（如政策、法规、补贴、价格、活动、行业动态等），必须结合【联网检索结果】核验，并在回答中说明信息来源与日期；若联网结果与公司制度冲突，提示以公司内部最新通知为准。
3. 若内部资料库与联网结果都不足以回答，明确说明「暂未找到确切依据」，不要编造。
4. 回答要简洁、口语化、对新人友好，能分点给出操作步骤的尽量分点。
5. 回答末尾用「📚 参考来源」列出引用（文档名 / 网页标题与链接）。
6. 涉及具体流程（办证、拓店登记等）要给出清晰步骤。
7. 排版紧凑：Markdown 段落之间不要插空行，列表项之间只用一个换行；用 # / ## 区分层级即可，不要多余空行拉大间距。手机屏幕阅读，保持视觉紧凑。
8. **同类枚举项单行输出**：配料、清单、要点等「并列性质」的项，**不要每项一行列成列表**，应直接用「、」或「，」连成一行（如：猪肉、花胶、五指毛桃粉、木薯淀粉、食用盐、白砂糖）。
9. **需要分维度说明时，用「小标题：内容」的紧凑单行写法，不要使用 - / * 列表块**：例如写「**主料**：客家猪肉（含量90%以上，选后腿肉）。**辅料**：木薯淀粉、食用盐、白砂糖等基础调味，不添加香精、色素、防腐剂。」而不是把它们逐行列成 bullet。仅以下情况才用真正的编号/列表：(a) 步骤型操作流程（如办证、拓店登记）需逐条执行；(b) 各项需要独立强调或排序。**默认整段紧凑自然语言，少用列表块**。"""

def build_messages(q, kb_chunks, web_results, assistant_name):
    kb_block = "\n\n".join(
        f"【资料片段 {i+1}｜来源：{c['source']}】\n{c['text']}"
        for i, c in enumerate(kb_chunks)
    ) or "（内部资料库未检索到相关内容）"
    web_block = "\n\n".join(
        f"【网页 {i+1}｜{w['title']}】\n{w['snippet']}\n链接：{w['url']}"
        for i, w in enumerate(web_results)
    ) if web_results else "（未触发联网核验 / 联网未返回结果）"
    user = f"""# 内部资料库检索结果
{kb_block}

# 联网核验结果
{web_block}

# 新员工提问
{q}

请综合以上信息作答，遵守系统设定中的回答规则。"""
    return [
        {"role": "system", "content": SYSTEM_PROMPT_TEMPLATE.format(name=assistant_name)},
        {"role": "user", "content": user},
    ]

# ---- DeepSeek 流式调用 ----
def stream_deepseek(messages):
    if not DEEPSEEK_API_KEY:
        return None
    try:
        r = requests.post(
            f"{DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                     "Content-Type": "application/json"},
            json={"model": DEEPSEEK_MODEL, "messages": messages, "stream": True,
                  "temperature": 0.3},
            timeout=60, stream=True,
        )
        if r.status_code != 200:
            print(f"[DEEPSEEK] {r.status_code} {r.text[:200]}")
            return None
        for line in r.iter_lines(decode_unicode=True):
            if not line or not line.startswith("data:"):
                continue
            data = line[5:].strip()
            if data == "[DONE]":
                break
            try:
                obj = json.loads(data)
                delta = obj["choices"][0]["delta"].get("content", "")
                if delta:
                    yield delta
            except Exception:
                continue
    except Exception as e:
        print(f"[DEEPSEEK] 异常: {e}")
        return None

def deepseek_json(messages, max_tokens=6000):
    """非流式调用 DeepSeek，要求返回 JSON 并解析（用于后台资料解析/生成题目）。"""
    if not DEEPSEEK_API_KEY:
        return None, "未配置 DEEPSEEK_API_KEY"
    try:
        r = requests.post(
            f"{DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                     "Content-Type": "application/json"},
            json={"model": DEEPSEEK_MODEL, "messages": messages,
                  "temperature": 0.2, "max_tokens": max_tokens,
                  "response_format": {"type": "json_object"}},
            timeout=240,
        )
        if r.status_code != 200:
            return None, f"DeepSeek 返回 {r.status_code}：{r.text[:200]}"
        text = r.json()["choices"][0]["message"]["content"]
        return json.loads(text), None
    except json.JSONDecodeError:
        return None, "模型输出不是合法 JSON，请重试"
    except Exception as e:
        return None, f"DeepSeek 调用异常：{e}"

def fallback_answer(q, kb_chunks, web_results):
    parts = ["（演示模式：未配置 DeepSeek API Key，以下为检索到的内部资料片段）\n"]
    if kb_chunks:
        parts.append("📚 内部资料库相关片段：")
        for c in kb_chunks:
            parts.append(f"\n— 来源《{c['source']}》—\n{c['text']}")
    if web_results:
        parts.append("\n🌐 联网核验结果：")
        for w in web_results:
            parts.append(f"\n· {w['title']}\n  {w['snippet']}\n  {w['url']}")
    if not kb_chunks and not web_results:
        parts.append("未检索到相关内容，请换一种问法或联系培训负责人。")
    return "\n".join(parts)

# ---- SSE 工具 ----
def sse(event, data):
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

# ============ 账号体系 / 后台 ============
def public_user(u):
    return {"id": u.id, "phone": u.phone, "name": u.name, "role": u.role,
            "hire_date": u.hire_date.isoformat() if u.hire_date else None,
            "position": u.position or "", "exam_opened": bool(u.exam_opened),
            "created_at": u.created_at.isoformat() if u.created_at else None,
            "last_login": u.last_login.isoformat() if u.last_login else None}

def get_setting(key, default=""):
    s = db.session.get(Setting, key)
    return s.value if s else default

def set_setting(key, value):
    s = db.session.get(Setting, key)
    if s:
        s.value = value
    else:
        db.session.add(Setting(key=key, value=value))
    db.session.commit()

def init_settings():
    for k, v in [("assistant_name", "喜客丸AI小助手"), ("announcement", "")]:
        if not db.session.get(Setting, k):
            db.session.add(Setting(key=k, value=v))
    db.session.commit()

def seed_admin():
    # 主账号始终以环境变量为准：不存在则创建，已存在则同步角色与密码。
    # 这样改 ADMIN_PHONE/ADMIN_DEFAULT_PASSWORD + 重新部署即可随时修正主账号。
    if not (ADMIN_PHONE and ADMIN_DEFAULT_PASSWORD):
        return
    u = User.query.filter_by(phone=ADMIN_PHONE).first()
    if not u:
        u = User(phone=ADMIN_PHONE, name="总部主账号", role="admin")
        db.session.add(u)
        print(f"[AUTH] 已创建总部主账号 {ADMIN_PHONE}")
    else:
        print(f"[AUTH] 主账号已存在，同步权限与密码 {ADMIN_PHONE}")
    u.role = "admin"
    u.name = u.name or "总部主账号"
    u.password_hash = generate_password_hash(ADMIN_DEFAULT_PASSWORD)
    db.session.commit()

def admin_required(f):
    @wraps(f)
    def dec(*a, **k):
        uid = session.get("user_id")
        if not uid:
            return jsonify({"error": "unauthorized"}), 401
        u = db.session.get(User, uid)
        if not u or u.role != "admin":
            return jsonify({"error": "forbidden"}), 403
        return f(*a, **k)
    return dec

# ============ Flask 应用 ============
app = Flask(__name__, static_folder=STATIC_DIR)
app.secret_key = SECRET_KEY
app.config["SQLALCHEMY_DATABASE_URI"] = make_db_uri()
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
# Neon/云 Postgres 会断开空闲连接；开启 pre-ping + 回收，避免命中死连接导致 500
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_pre_ping": True,
    "pool_recycle": 280,
}
db.init_app(app)

def _column_exists(table, column):
    """判断某表是否存在某列（兼容 SQLite / Postgres）。"""
    try:
        if db.engine.dialect.name == "sqlite":
            rows = db.session.execute(text(f"PRAGMA table_info({table})")).fetchall()
            return any(r[1] == column for r in rows)
        rows = db.session.execute(
            text("SELECT column_name FROM information_schema.columns "
                 "WHERE table_name=:t AND column_name=:c"),
            {"t": table, "c": column}).fetchall()
        return len(rows) > 0
    except Exception:
        return False

def migrate_db():
    """建表 + 为已存在的表补加新列（create_all 不会 alter 已有表）。"""
    db.create_all()
    alters = [
        ("users", "hire_date", "DATE"),
        ("users", "position", "VARCHAR(50)"),
        ("users", "exam_opened", "BOOLEAN DEFAULT FALSE"),
    ]
    for table, col, sql in alters:
        if not _column_exists(table, col):
            try:
                db.session.execute(text(f"ALTER TABLE {table} ADD COLUMN {col} {sql}"))
                db.session.commit()
                print(f"[MIGRATE] 已为 {table} 增加列 {col}")
            except Exception as e:
                db.session.rollback()
                print(f"[MIGRATE] 增加列 {table}.{col} 失败: {e}")

with app.app_context():
    migrate_db()
    init_settings()
    seed_admin()

# 启动时预加载知识库（gunicorn 多 worker 下每个进程各自加载）
try:
    load_kb()
    load_onboarding()
    load_quiz()
except Exception as e:
    print(f"[KB] 启动加载失败（将在首次提问时重试）: {e}")


@app.route("/")
def index():
    return send_from_directory(STATIC_DIR, "index.html")


@app.route("/admin")
def admin_page():
    uid = session.get("user_id")
    if not uid:
        return redirect("/")
    u = db.session.get(User, uid)
    if not u or u.role != "admin":
        return "无权限访问（需总部主账号）", 403
    return render_template("admin.html")


@app.route("/api/health")
def health():
    return json.dumps({"status": "ok", "kb_chunks": len(_chunks),
                       "web": WEB_SEARCH_ENABLED, "key": bool(DEEPSEEK_API_KEY)})


@app.route("/api/reload", methods=["POST"])
def reload_kb():
    try:
        load_kb()
        return json.dumps({"status": "ok", "kb_chunks": len(_chunks)})
    except Exception as e:
        return json.dumps({"status": "error", "message": str(e)})


# ---- 账号：注册 / 登录 / 登出 / 当前用户 ----
@app.route("/api/register", methods=["POST"])
def register():
    data = request.get_json(silent=True) or {}
    phone = (data.get("phone") or "").strip()
    pwd = (data.get("password") or "")
    name = (data.get("name") or "").strip() or None
    if not re.match(r"^1[3-9]\d{9}$", phone):
        return jsonify({"error": "手机号格式不正确（需 11 位大陆手机号）"}), 400
    if len(pwd) < 6:
        return jsonify({"error": "密码至少 6 位"}), 400
    if User.query.filter_by(phone=phone).first():
        return jsonify({"error": "该手机号已注册，请直接登录"}), 400
    role = "admin" if (ADMIN_PHONE and phone == ADMIN_PHONE) else "employee"
    u = User(phone=phone, password_hash=generate_password_hash(pwd), name=name, role=role)
    db.session.add(u)
    db.session.commit()
    session["user_id"] = u.id
    session["role"] = u.role
    return jsonify({"ok": True, "user": public_user(u)})


@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json(silent=True) or {}
    phone = (data.get("phone") or "").strip()
    pwd = (data.get("password") or "")
    u = User.query.filter_by(phone=phone).first()
    if not u or not check_password_hash(u.password_hash, pwd):
        return jsonify({"error": "手机号或密码错误"}), 401
    u.last_login = datetime.utcnow()
    db.session.commit()
    session["user_id"] = u.id
    session["role"] = u.role
    return jsonify({"ok": True, "user": public_user(u)})


@app.route("/api/logout", methods=["POST"])
def logout():
    session.clear()
    return jsonify({"ok": True})


@app.route("/api/me")
def me():
    uid = session.get("user_id")
    if not uid:
        return jsonify({"user": None})
    u = db.session.get(User, uid)
    if not u:
        return jsonify({"user": None})
    return jsonify({"user": public_user(u)})


# ---- 配置（员工端拉取：动态助手名 + 群公告）----
@app.route("/api/config")
def config():
    def _read():
        return jsonify({
            "assistant_name": get_setting("assistant_name", "喜客丸AI小助手"),
            "announcement": get_setting("announcement", ""),
        })
    try:
        return _read()
    except Exception as e:
        # 云 Postgres 偶发连接抖动：回滚后重试一次，仍失败则返回安全默认值
        print(f"[CONFIG] 首次读取失败，重试: {e}")
        db.session.rollback()
        try:
            return _read()
        except Exception as e2:
            print(f"[CONFIG] 读取配置失败: {e2}")
            return jsonify({
                "assistant_name": "喜客丸AI小助手",
                "announcement": "",
            })


# ---- 找回密码：提交重置申请，由总部主账号在后台处理 ----
@app.route("/api/forgot", methods=["POST"])
def forgot():
    data = request.get_json(silent=True) or {}
    phone = (data.get("phone") or "").strip()
    u = User.query.filter_by(phone=phone).first()
    if not u:
        return jsonify({"error": "该手机号未注册"}), 404
    db.session.add(ResetRequest(user_id=u.id, status="pending"))
    db.session.commit()
    return jsonify({"ok": True, "message": "已提交密码重置申请，请联系总部管理员处理"})


# ---- 问答（需登录，并埋点使用记录）----
@app.route("/api/ask", methods=["POST"])
def ask():
    uid = session.get("user_id")
    if not uid:
        return Response(sse("error", {"message": "请先登录后再使用"}),
                        mimetype="text/event-stream")
    data = request.get_json(silent=True) or {}
    q = (data.get("question") or "").strip()
    if not q:
        return Response(sse("error", {"message": "请输入问题"}), mimetype="text/event-stream")

    kb_chunks = retrieve(q, TOP_K)
    need_web = is_time_sensitive(q) or (not kb_chunks)
    web_results = web_search(q) if need_web else []
    assistant_name = get_setting("assistant_name", "喜客丸AI小助手")
    messages = build_messages(q, kb_chunks, web_results, assistant_name)

    def gen():
        yield sse("sources", {"kb": [c["source"] for c in kb_chunks],
                              "web": [w["title"] for w in web_results],
                              "web_used": bool(web_results)})
        acc = ""
        if DEEPSEEK_API_KEY:
            gen_ = stream_deepseek(messages)
            if gen_ is None:
                for piece in fallback_answer(q, kb_chunks, web_results).split("\n"):
                    acc += piece + "\n"
                    yield sse("token", piece + "\n")
            else:
                for delta in gen_:
                    acc += delta
                    yield sse("token", delta)
        else:
            for piece in fallback_answer(q, kb_chunks, web_results).split("\n"):
                acc += piece + "\n"
                yield sse("token", piece + "\n")
        # 埋点：记录本次提问（含回答预览）。
        # SSE 生成器在请求上下文之外执行，db 提交需显式绑定应用上下文。
        log_id = None
        try:
            with app.app_context():
                log = UsageLog(user_id=uid, question=q,
                               answer_preview=acc[:200],
                               kb_count=len(kb_chunks),
                               web_count=len(web_results))
                db.session.add(log)
                db.session.commit()
                log_id = log.id
        except Exception as e:
            print(f"[USAGE] 写入失败: {e}")
        yield sse("done", {"kb_count": len(kb_chunks), "web_count": len(web_results), "log_id": log_id})

    return Response(gen(), mimetype="text/event-stream")


# ============ 新员工必读 + 考试（员工端，需登录）============
def _require_uid():
    uid = session.get("user_id")
    return uid

def _onb_status(uid):
    progresses = {lp.module_id: lp for lp in LearningProgress.query.filter_by(user_id=uid).all()}
    modules = []
    all_done = True
    for m in _onb_modules:
        lp = progresses.get(m["id"])
        done = bool(lp and lp.completed)
        if not done:
            all_done = False
        modules.append({"id": m["id"], "title": m["title"], "completed": done})
    u = db.session.get(User, uid)
    exam_opened = bool(u and u.exam_opened)
    exam_available = exam_opened and all_done
    return modules, all_done, exam_opened, exam_available

@app.route("/api/onboarding/modules")
def onboarding_modules():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    out = [{"id": m["id"], "title": m["title"], "summary": m.get("summary", ""),
            "content": m.get("content", [])} for m in _onb_modules]
    return jsonify({"modules": out})

@app.route("/api/onboarding/status")
def onboarding_status():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    modules, all_done, exam_opened, exam_available = _onb_status(uid)
    return jsonify({"modules": modules, "all_completed": all_done,
                    "exam_opened": exam_opened, "exam_available": exam_available})

@app.route("/api/onboarding/mark", methods=["POST"])
def onboarding_mark():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    data = request.get_json(silent=True) or {}
    mid = (data.get("module_id") or "").strip()
    completed = bool(data.get("completed", True))
    if not mid or not any(m["id"] == mid for m in _onb_modules):
        return jsonify({"error": "无效模块"}), 400
    lp = LearningProgress.query.filter_by(user_id=uid, module_id=mid).first()
    if not lp:
        lp = LearningProgress(user_id=uid, module_id=mid)
        db.session.add(lp)
    lp.completed = completed
    lp.completed_at = datetime.utcnow() if completed else None
    db.session.commit()
    return jsonify({"ok": True, "completed": completed})


def _score_attempt(attempt):
    qids = json.loads(attempt.question_ids)
    answers = json.loads(attempt.answers)
    score = 0
    for qid in qids:
        q = quiz_question_by_id(qid)
        if not q:
            continue
        if answers.get(str(qid), "") == q.get("answer", ""):
            score += QUIZ_PER_SCORE
    attempt.score = score
    attempt.total = len(qids) * QUIZ_PER_SCORE
    attempt.status = "submitted"

def _auto_settle(attempt):
    if attempt and attempt.status == "in_progress" and attempt.deadline \
            and datetime.utcnow() >= attempt.deadline:
        _score_attempt(attempt)
        db.session.commit()
    return attempt

@app.route("/api/exam/state")
def exam_state():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    attempt = (QuizAttempt.query.filter_by(user_id=uid)
               .order_by(QuizAttempt.created_at.desc()).first())
    if not attempt:
        return jsonify({"status": "none"})
    _auto_settle(attempt)
    qids = json.loads(attempt.question_ids)
    answers = json.loads(attempt.answers)
    if attempt.status == "in_progress":
        questions = [quiz_to_front(q, False) for qid in qids
                     if (q := quiz_question_by_id(qid))]
        remaining = int((attempt.deadline - datetime.utcnow()).total_seconds()) if attempt.deadline else 0
        return jsonify({"status": "in_progress", "deadline": attempt.deadline.isoformat(),
                        "remaining": max(0, remaining), "questions": questions, "answers": answers})
    details = []
    for qid in qids:
        q = quiz_question_by_id(qid)
        if q:
            d = quiz_to_front(q, True)
            d["selected"] = answers.get(str(qid), "")
            details.append(d)
    return jsonify({"status": "submitted", "score": attempt.score, "total": attempt.total,
                    "submitted_at": attempt.created_at.isoformat() if attempt.created_at else None,
                    "details": details})

@app.route("/api/exam/start", methods=["POST"])
def exam_start():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    modules, all_done, exam_opened, exam_available = _onb_status(uid)
    if not exam_available:
        return jsonify({"error": "考试尚未开启，或必读板块未全部学完"}), 400
    existing = (QuizAttempt.query.filter_by(user_id=uid)
                .order_by(QuizAttempt.created_at.desc()).first())
    if existing and existing.status == "in_progress":
        return jsonify({"error": "已有进行中的考试，请继续作答"}), 400
    if existing and existing.status == "submitted":
        return jsonify({"error": "已完成考试，可在结果页查看"}), 400
    now = datetime.utcnow()
    qids = draw_questions(20)
    attempt = QuizAttempt(user_id=uid, question_ids=json.dumps(qids),
                          answers=json.dumps({}), started_at=now,
                          deadline=now + timedelta(minutes=QUIZ_DURATION),
                          duration_min=QUIZ_DURATION, status="in_progress")
    db.session.add(attempt)
    db.session.commit()
    questions = [quiz_to_front(q, False) for qid in qids
                 if (q := quiz_question_by_id(qid))]
    return jsonify({"status": "in_progress", "deadline": attempt.deadline.isoformat(),
                    "remaining": QUIZ_DURATION * 60, "questions": questions, "answers": {}})

@app.route("/api/exam/answer", methods=["POST"])
def exam_answer():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    data = request.get_json(silent=True) or {}
    qid = data.get("question_id")
    sel = data.get("selected", "")
    attempt = (QuizAttempt.query.filter_by(user_id=uid, status="in_progress")
               .order_by(QuizAttempt.created_at.desc()).first())
    if not attempt:
        return jsonify({"error": "没有进行中的考试"}), 400
    if attempt.deadline and datetime.utcnow() >= attempt.deadline:
        _score_attempt(attempt)
        db.session.commit()
        return jsonify({"error": "考试时间已结束，已自动结算", "settled": True}), 400
    answers = json.loads(attempt.answers)
    answers[str(qid)] = sel
    attempt.answers = json.dumps(answers)
    db.session.commit()
    return jsonify({"ok": True})

@app.route("/api/exam/submit", methods=["POST"])
def exam_submit():
    uid = _require_uid()
    if not uid:
        return jsonify({"error": "unauthorized"}), 401
    attempt = (QuizAttempt.query.filter_by(user_id=uid, status="in_progress")
               .order_by(QuizAttempt.created_at.desc()).first())
    if not attempt:
        return jsonify({"error": "没有进行中的考试"}), 400
    _score_attempt(attempt)
    db.session.commit()
    qids = json.loads(attempt.question_ids)
    answers = json.loads(attempt.answers)
    details = []
    for qid in qids:
        q = quiz_question_by_id(qid)
        if q:
            d = quiz_to_front(q, True)
            d["selected"] = answers.get(str(qid), "")
            details.append(d)
    return jsonify({"status": "submitted", "score": attempt.score, "total": attempt.total,
                    "details": details})


# ============ 总部后台 API（均需 admin 角色）============
@app.route("/api/admin/stats")
@admin_required
def admin_stats():
    total_users = User.query.count()
    admins = User.query.filter_by(role="admin").count()
    total_q = UsageLog.query.count()
    today_start = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    today_q = UsageLog.query.filter(UsageLog.created_at >= today_start).count()
    return jsonify({"total_users": total_users, "admins": admins,
                    "total_questions": total_q, "questions_today": today_q,
                    "assistant_name": get_setting("assistant_name", "喜客丸AI小助手"),
                    "announcement": get_setting("announcement", "")})


@app.route("/api/admin/balance")
@admin_required
def admin_balance():
    if not DEEPSEEK_API_KEY:
        return jsonify({"configured": False})
    try:
        r = requests.get(DEEPSEEK_BASE_URL + "/user/balance",
                         headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, timeout=10)
        return jsonify({"configured": True, "data": r.json()})
    except Exception as e:
        return jsonify({"configured": True, "error": str(e)})


@app.route("/api/admin/announcement", methods=["POST"])
@admin_required
def admin_announcement():
    data = request.get_json(silent=True) or {}
    set_setting("announcement", data.get("text", "") or "")
    return jsonify({"ok": True})


@app.route("/api/admin/rename", methods=["POST"])
@admin_required
def admin_rename():
    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "名称不能为空"}), 400
    set_setting("assistant_name", name)
    return jsonify({"ok": True, "assistant_name": name})


@app.route("/api/admin/users")
@admin_required
def admin_users():
    q = (request.args.get("q") or "").strip()
    page = max(1, int(request.args.get("page", "1")))
    per = 50
    query = User.query
    if q:
        like = f"%{q}%"
        query = query.filter((User.phone.like(like)) | (User.name.like(like)))
    total = query.count()
    users = (query.order_by(User.created_at.desc())
             .offset((page - 1) * per).limit(per).all())
    items = []
    for u in users:
        cnt = UsageLog.query.filter_by(user_id=u.id).count()
        last = (UsageLog.query.filter_by(user_id=u.id)
                .order_by(UsageLog.created_at.desc()).first())
        items.append({**public_user(u), "usage_count": cnt,
                      "last_question": last.question if last else None,
                      "last_active": last.created_at.isoformat() if last else None,
                      "exam_opened": bool(u.exam_opened), "position": u.position or ""})
    return jsonify({"total": total, "page": page, "items": items})


@app.route("/api/admin/users/<int:uid>/role", methods=["POST"])
@admin_required
def admin_role(uid):
    data = request.get_json(silent=True) or {}
    role = data.get("role")
    if role not in ("admin", "employee"):
        return jsonify({"error": "无效角色"}), 400
    u = User.query.get_or_404(uid)
    u.role = role
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/admin/users/<int:uid>", methods=["DELETE"])
@admin_required
def admin_delete_user(uid):
    me_id = session.get("user_id")
    target = db.session.get(User, uid)
    if not target:
        return jsonify({"error": "员工不存在"}), 404
    if target.id == me_id:
        return jsonify({"error": "不能删除自己的主账号"}), 400
    if target.role == "admin":
        other_admins = User.query.filter(User.role == "admin", User.id != uid).count()
        if other_admins == 0:
            return jsonify({"error": "至少保留一个管理员，无法删除最后一位管理员"}), 400
    # 级联清理（顺序：先删依赖 UsageLog 的 Feedback，再删其余子表，最后删用户）
    # Postgres 严格外键约束，漏清任一子表都会 500；SQLite 本地默认不强制外键故难复现
    Feedback.query.filter_by(user_id=uid).delete()
    LearningProgress.query.filter_by(user_id=uid).delete()
    QuizAttempt.query.filter_by(user_id=uid).delete()
    UsageLog.query.filter_by(user_id=uid).delete()
    ResetRequest.query.filter_by(user_id=uid).delete()
    phone = target.phone
    db.session.delete(target)
    db.session.commit()
    return jsonify({"ok": True, "deleted": phone})


@app.route("/api/admin/reset-password", methods=["POST"])
@admin_required
def admin_reset():
    data = request.get_json(silent=True) or {}
    uid = data.get("user_id")
    pwd = (data.get("new_password") or "")
    if len(pwd) < 6:
        return jsonify({"error": "新密码至少 6 位"}), 400
    u = db.session.get(User, uid)
    if not u:
        return jsonify({"error": "用户不存在"}), 404
    u.password_hash = generate_password_hash(pwd)
    ResetRequest.query.filter_by(user_id=uid, status="pending").update(
        {ResetRequest.status: "done"})
    db.session.commit()
    return jsonify({"ok": True})


# ---- 总部代建员工账号 ----
@app.route("/api/admin/users", methods=["POST"])
@admin_required
def admin_create_user():
    data = request.get_json(silent=True) or {}
    phone = (data.get("phone") or "").strip()
    pwd = (data.get("password") or "")
    name = (data.get("name") or "").strip() or None
    position = (data.get("position") or "").strip() or None
    hire_date = data.get("hire_date") or None
    if not re.match(r"^1[3-9]\d{9}$", phone):
        return jsonify({"error": "手机号格式不正确（需 11 位大陆手机号）"}), 400
    if len(pwd) < 6:
        return jsonify({"error": "初始密码至少 6 位"}), 400
    if User.query.filter_by(phone=phone).first():
        return jsonify({"error": "该手机号已存在"}), 400
    hd = None
    if hire_date:
        try:
            hd = datetime.strptime(hire_date, "%Y-%m-%d").date()
        except Exception:
            return jsonify({"error": "入职日期格式应为 YYYY-MM-DD"}), 400
    u = User(phone=phone, password_hash=generate_password_hash(pwd), name=name,
             role="employee", position=position, hire_date=hd, exam_opened=False)
    db.session.add(u)
    db.session.commit()
    return jsonify({"ok": True, "user": public_user(u)})


# ---- 总部为某员工开启考试 ----
@app.route("/api/admin/exam/open", methods=["POST"])
@admin_required
def admin_exam_open():
    data = request.get_json(silent=True) or {}
    uid = data.get("user_id")
    u = db.session.get(User, uid)
    if not u:
        return jsonify({"error": "员工不存在"}), 404
    # 清理该员工已有的考试尝试，使其可重新参加（若之前考过）
    QuizAttempt.query.filter_by(user_id=u.id).delete()
    db.session.commit()
    u.exam_opened = True
    db.session.commit()
    return jsonify({"ok": True, "exam_opened": True})


# ---- 总部后台：考试结果总览 + 易错题统计 ----
@app.route("/api/admin/quiz/results")
@admin_required
def admin_quiz_results():
    attempts = (QuizAttempt.query.filter_by(status="submitted")
                .order_by(QuizAttempt.created_at.desc()).all())
    items = []
    # 易错题统计
    wrong_counter = {}
    for a in attempts:
        u = db.session.get(User, a.user_id)
        qids = json.loads(a.question_ids)
        answers = json.loads(a.answers)
        wrong_ids = []
        for qid in qids:
            q = quiz_question_by_id(qid)
            if not q:
                continue
            if answers.get(str(qid), "") != q.get("answer", ""):
                wrong_ids.append(qid)
                wrong_counter[qid] = wrong_counter.get(qid, 0) + 1
        items.append({
            "attempt_id": a.id, "user_id": a.user_id,
            "phone": u.phone if u else "-", "name": u.name if u else "-",
            "score": a.score, "total": a.total,
            "submitted_at": a.created_at.isoformat() if a.created_at else None,
            "wrong_count": len(wrong_ids),
        })
    ranked = sorted(wrong_counter.items(), key=lambda kv: kv[1], reverse=True)
    weak = []
    for qid, cnt in ranked[:10]:
        q = quiz_question_by_id(qid)
        if q:
            weak.append({"id": qid, "stem": q["stem"], "topic": q.get("topic", ""),
                         "wrong_count": cnt, "answer": q.get("answer", "")})
    return jsonify({"items": items, "count": len(items), "weak": weak})


# ---- 后台员工详情：补充必读进度 + 考试 ----
def _detail_learning(uid):
    progresses = {lp.module_id: lp for lp in LearningProgress.query.filter_by(user_id=uid).all()}
    modules = [{"id": m["id"], "title": m["title"], "completed": bool(lp and lp.completed)}
               for m in _onb_modules
               for lp in [progresses.get(m["id"])]]
    all_done = all(m["completed"] for m in modules)
    return modules, all_done

def _detail_exam(uid):
    attempt = (QuizAttempt.query.filter_by(user_id=uid)
               .order_by(QuizAttempt.created_at.desc()).first())
    if not attempt:
        return None
    qids = json.loads(attempt.question_ids)
    answers = json.loads(attempt.answers)
    details = []
    for qid in qids:
        q = quiz_question_by_id(qid)
        if q:
            d = quiz_to_front(q, True)
            d["selected"] = answers.get(str(qid), "")
            details.append(d)
    return {"status": attempt.status, "score": attempt.score, "total": attempt.total,
            "submitted_at": attempt.created_at.isoformat() if attempt.created_at else None,
            "details": details}


@app.route("/api/admin/users/<int:uid>")
@admin_required
def admin_user_detail(uid):
    u = User.query.get_or_404(uid)
    logs = (UsageLog.query.filter_by(user_id=uid)
            .order_by(UsageLog.created_at.desc()).limit(100).all())
    modules, all_done = _detail_learning(uid)
    exam = _detail_exam(uid)
    return jsonify({"user": public_user(u),
                    "learning": {"modules": modules, "all_completed": all_done},
                    "exam": exam,
                    "usage": [{"question": l.question, "preview": l.answer_preview,
                               "kb": l.kb_count, "web": l.web_count,
                               "at": l.created_at.isoformat()} for l in logs]})


# ---- 员工端：答案反馈 👍/👎 ----
@app.route("/api/feedback", methods=["POST"])
def submit_feedback():
    uid = session.get("user_id")
    if not uid:
        return jsonify({"error": "请先登录"}), 401
    data = request.get_json(silent=True) or {}
    log_id = data.get("log_id")
    rating = data.get("rating")  # "up" | "down"
    comment = (data.get("comment") or "").strip() or None
    if not log_id or rating not in ("up", "down"):
        return jsonify({"error": "参数无效"}), 400
    log = db.session.get(UsageLog, log_id)
    if not log or log.user_id != uid:
        return jsonify({"error": "记录不存在或无权操作"}), 404
    existing = Feedback.query.filter_by(usage_log_id=log_id, user_id=uid).first()
    if existing:
        existing.rating = rating
        existing.comment = comment
    else:
        db.session.add(Feedback(usage_log_id=log_id, user_id=uid,
                                 rating=rating, comment=comment))
    db.session.commit()
    return jsonify({"ok": True})


# ---- 后台：热门问题榜（近 30 天 Top 10）----
@app.route("/api/admin/hot-questions")
@admin_required
def admin_hot_questions():
    since = datetime.utcnow() - timedelta(days=30)
    cnt = db.func.count(UsageLog.id)
    rows = (db.session.query(UsageLog.question, cnt.label("c"))
            .filter(UsageLog.created_at >= since)
            .group_by(UsageLog.question)
            .order_by(cnt.desc())
            .limit(10).all())
    return jsonify({"items": [{"question": r.question, "count": r.c} for r in rows]})


# ---- 后台：反馈管理列表 ----
@app.route("/api/admin/feedback")
@admin_required
def admin_feedback_list():
    rating = request.args.get("rating", "")
    per = 50
    query = Feedback.query
    if rating in ("up", "down"):
        query = query.filter(Feedback.rating == rating)
    total = query.count()
    items = (query.order_by(Feedback.created_at.desc()).limit(per).all())
    result = []
    for f in items:
        log = db.session.get(UsageLog, f.usage_log_id) if f.usage_log_id else None
        u = db.session.get(User, f.user_id) if f.user_id else None
        result.append({
            "id": f.id, "rating": f.rating, "comment": f.comment or "",
            "created_at": f.created_at.isoformat() if f.created_at else None,
            "user_phone": u.phone if u else "-",
            "user_name": u.name if u else "-",
            "question": log.question if log else "-",
            "answer_preview": (log.answer_preview[:120] + "...") if log and log.answer_preview and len(log.answer_preview) > 120 else (log.answer_preview or "" if log else ""),
        })
    return jsonify({"total": total, "items": result})


# ---- 后台：在线管理知识库 ----
_KB_NAME_RE = re.compile(r"^[\w\-]+\.txt$")


@app.route("/api/admin/kb/files")
@admin_required
def admin_kb_files():
    files = []
    if os.path.isdir(KB_DIR):
        for f in sorted(os.listdir(KB_DIR)):
            if f.endswith(".txt"):
                path = os.path.join(KB_DIR, f)
                size = os.path.getsize(path)
                source = os.path.splitext(f)[0]
                chunk_count = sum(1 for c in _chunks if c.get("source") == source)
                files.append({"name": f, "size": size, "chunks": chunk_count})
    return jsonify({"items": files, "total_chunks": len(_chunks)})


@app.route("/api/admin/kb/file/<filename>")
@admin_required
def admin_kb_read(filename):
    if not _KB_NAME_RE.match(filename):
        return jsonify({"error": "文件名不合法（仅允许字母数字下划线横线.txt）"}), 400
    path = os.path.join(KB_DIR, filename)
    if not os.path.exists(path):
        return jsonify({"error": "文件不存在"}), 404
    with open(path, encoding="utf-8", errors="ignore") as fh:
        content = fh.read()
    return jsonify({"name": filename, "content": content})


@app.route("/api/admin/kb/file/<filename>", methods=["POST"])
@admin_required
def admin_kb_save(filename):
    if not _KB_NAME_RE.match(filename):
        return jsonify({"error": "文件名不合法（仅允许字母数字下划线横线.txt）"}), 400
    data = request.get_json(silent=True) or {}
    content = data.get("content", "")
    os.makedirs(KB_DIR, exist_ok=True)
    path = os.path.join(KB_DIR, filename)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(content)
    return jsonify({"ok": True, "name": filename})


@app.route("/api/admin/kb/file/<filename>", methods=["DELETE"])
@admin_required
def admin_kb_delete(filename):
    if not _KB_NAME_RE.match(filename):
        return jsonify({"error": "文件名不合法"}), 400
    path = os.path.join(KB_DIR, filename)
    if not os.path.exists(path):
        return jsonify({"error": "文件不存在"}), 404
    os.remove(path)
    return jsonify({"ok": True, "deleted": filename})


@app.route("/api/admin/kb/reload", methods=["POST"])
@admin_required
def admin_kb_reload():
    try:
        load_kb()
        return jsonify({"ok": True, "kb_chunks": len(_chunks)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============ 后台：必读板块 / 题库在线编辑 + 资料上传AI解析 ============
QUIZ_TOPICS = ["产品知识", "抖音来客平台", "门店准入", "入驻流程", "品牌授权", "入驻审核"]

def _validate_onb_doc(doc):
    """校验必读板块文档结构，返回 (规范后文档, 错误信息)。"""
    if not isinstance(doc, dict) or not isinstance(doc.get("modules"), list):
        return None, "格式错误：缺少 modules 数组"
    if not doc["modules"]:
        return None, "至少保留 1 个必读板块"
    seen = set()
    for i, m in enumerate(doc["modules"]):
        mid = str(m.get("id", "")).strip()
        title = str(m.get("title", "")).strip()
        if not mid or not title:
            return None, f"第 {i+1} 个板块缺少 id 或标题"
        if mid in seen:
            return None, f"板块 id 重复：{mid}"
        seen.add(mid)
        if not isinstance(m.get("content"), list):
            return None, f"板块「{title}」的 content 必须是数组（每行一条）"
    return doc, None

def _validate_quiz_doc(doc):
    """校验题库文档结构，返回 (规范后文档, 错误信息)。
    题目 schema：{id, topic, stem, options:{A:...,B:...}, answer, explanation?}"""
    if not isinstance(doc, dict) or not isinstance(doc.get("questions"), list):
        return None, "格式错误：缺少 questions 数组"
    if not doc["questions"]:
        return None, "至少保留 1 道题"
    seen = set()
    for i, q in enumerate(doc["questions"]):
        stem = str(q.get("stem", "")).strip()
        opts = q.get("options")
        ans = str(q.get("answer", "")).strip().upper()
        topic = str(q.get("topic", "")).strip()
        if not stem:
            return None, f"第 {i+1} 题缺少题干"
        if not isinstance(opts, dict) or len(opts) < 2:
            return None, f"「{stem[:15]}…」选项至少 2 个"
        if ans not in opts:
            return None, f"「{stem[:15]}…」答案必须是 {'/'.join(opts.keys())}"
        if not topic:
            return None, f"「{stem[:15]}…」缺少主题"
        qid = q.get("id")
        if qid in seen:
            return None, f"题目 id 重复：{qid}"
        seen.add(qid)
    return doc, None

def _save_onb_doc(doc):
    _save_setting(ONB_DB_KEY, json.dumps(doc, ensure_ascii=False))
    load_onboarding()

def _save_quiz_doc(doc):
    _save_setting(QUIZ_DB_KEY, json.dumps(doc, ensure_ascii=False))
    load_quiz()


@app.route("/api/admin/onboarding")
@admin_required
def admin_onboarding_get():
    return jsonify({"doc": _onb_doc, "source": _onb_source})


@app.route("/api/admin/onboarding", methods=["POST"])
@admin_required
def admin_onboarding_save():
    data = request.get_json(silent=True) or {}
    doc = data.get("doc")
    doc, err = _validate_onb_doc(doc)
    if err:
        return jsonify({"error": err}), 400
    try:
        _save_onb_doc(doc)
        return jsonify({"ok": True, "modules": len(_onb_modules), "source": _onb_source})
    except Exception as e:
        return jsonify({"error": f"保存失败：{e}"}), 500


@app.route("/api/admin/quiz")
@admin_required
def admin_quiz_get():
    return jsonify({"doc": _quiz_doc, "source": _quiz_source})


@app.route("/api/admin/quiz", methods=["POST"])
@admin_required
def admin_quiz_save():
    data = request.get_json(silent=True) or {}
    doc = data.get("doc")
    doc, err = _validate_quiz_doc(doc)
    if err:
        return jsonify({"error": err}), 400
    try:
        _save_quiz_doc(doc)
        return jsonify({"ok": True, "questions": len(_quiz_bank), "source": _quiz_source})
    except Exception as e:
        return jsonify({"error": f"保存失败：{e}"}), 500


# ---- 资料文件解析（后台上传 → 纯文本） ----
def _extract_file_text(stream, filename):
    """解析上传的表格/文档为纯文本。返回 (文本, 错误)。"""
    import io
    name = (filename or "").lower()
    try:
        raw = stream.read()
        if name.endswith((".txt", ".md")):
            for enc in ("utf-8", "gbk", "utf-16"):
                try:
                    return raw.decode(enc), None
                except Exception:
                    continue
            return None, "无法识别文件编码"
        if name.endswith(".csv"):
            import csv as _csv
            for enc in ("utf-8-sig", "gbk"):
                try:
                    text = raw.decode(enc)
                    break
                except Exception:
                    text = None
            if text is None:
                return None, "CSV 编码无法识别"
            rows = list(_csv.reader(io.StringIO(text)))
            lines = [" | ".join(str(c).strip() for c in row) for row in rows
                     if any(str(c).strip() for c in row)]
            return "\n".join(lines), None
        if name.endswith((".xlsx", ".xlsm")):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
            out = []
            for sn in wb.sheetnames:
                ws = wb[sn]
                out.append(f"【工作表：{sn}】")
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c).strip() for c in row]
                    if any(cells):
                        out.append(" | ".join(cells))
            return "\n".join(out), None
        if name.endswith(".docx"):
            from docx import Document
            d = Document(io.BytesIO(raw))
            out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            for tb in d.tables:
                for row in tb.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        out.append(" | ".join(cells))
            return "\n".join(out), None
        if name.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except ImportError:
                return None, "暂不支持 PDF，请先转为 Word/Excel/TXT 再上传"
            reader = PdfReader(io.BytesIO(raw))
            out = []
            for pg in reader.pages:
                t = (pg.extract_text() or "").strip()
                if t:
                    out.append(t)
            return "\n".join(out), None
        return None, f"暂不支持该格式（{filename}），支持：xlsx / csv / docx / pdf / txt / md"
    except Exception as e:
        return None, f"解析失败：{e}"


@app.route("/api/admin/material/analyze", methods=["POST"])
@admin_required
def admin_material_analyze():
    """上传资料文件 → 解析文本 → AI 生成「必读新增内容 + 新题目」建议（不落库，先预览）。"""
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"error": "请选择要上传的文件"}), 400
    text, err = _extract_file_text(f.stream, f.filename)
    if err:
        return jsonify({"error": err}), 400
    text = (text or "").strip()
    if len(text) < 30:
        return jsonify({"error": "解析到的内容太少（不足 30 字），请检查文件"}), 400
    material = text[:15000]
    # 现有内容摘要给模型，避免生成重复内容
    mod_brief = "\n".join(
        f"- 板块「{m['id']}｜{m['title']}」共 {len(m.get('content', []))} 条：" +
        "；".join(c[:40] for c in m.get("content", [])[:3]) + "…"
        for m in _onb_modules)
    q_brief = "\n".join(f"- [{q.get('topic')}] {q.get('question', '')[:40]}"
                        for q in _quiz_bank)
    sys_prompt = (
        "你是喜客丸（养生肉丸品牌）拓店岗培训内容运营。公司上传了一份新培训资料，"
        "请基于资料内容，为「新员工必读板块」和「入职考试题库」生成增量更新建议。\n"
        "要求：\n"
        "1. 只提取资料中【新的、对拓店新人有价值】的信息；与现有内容重复的不要生成。\n"
        "2. module_updates：按资料内容归入最合适的现有板块（用现有 module_id）；"
        "每条 new_item 是一行完整、可直接给新人阅读的知识条目，以【小标题】开头，简洁具体。\n"
        "3. new_questions：生成 3~6 道单项选择题，考资料中的关键知识点；"
        "topic 必须从 [" + "、".join(QUIZ_TOPICS) + "] 中选择；"
        "options 为 3~4 个选项文本（不含字母前缀）；answer 是正确选项字母；"
        "explanation 一句话说明。不要与现有题目重复。\n"
        "4. 若资料内容与培训无关（如聊天记录、无关表格），两项都返回空数组。\n"
        '输出 JSON：{"module_updates":[{"module_id":"...","new_items":["..."]}],'
        '"new_questions":[{"topic":"...","stem":"...","options":["..."],"answer":"B","explanation":"..."}]}'
    )
    user_prompt = (f"【现有必读板块】\n{mod_brief}\n\n【现有题库题目】\n{q_brief}\n\n"
                   f"【上传的资料《{f.filename}》解析文本】\n{material}")
    result, err = deepseek_json([
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt},
    ])
    if err:
        return jsonify({"error": err}), 500
    # 规范化 + 为新题分配 id
    mod_updates = []
    valid_ids = {m["id"] for m in _onb_modules}
    for mu in (result.get("module_updates") or []):
        mid = mu.get("module_id")
        items = [str(x).strip() for x in (mu.get("new_items") or []) if str(x).strip()]
        if mid in valid_ids and items:
            mod_updates.append({"module_id": mid, "new_items": items})
    next_id = max((q.get("id", 0) for q in _quiz_bank if isinstance(q.get("id"), int)),
                  default=0) + 1
    new_questions = []
    for q in (result.get("new_questions") or []):
        # AI 可能输出列表或字典两种 options 形式，统一归一化为 {A:...,B:...}
        raw_opts = q.get("options")
        if isinstance(raw_opts, dict):
            opts = [str(v).strip() for _, v in sorted(raw_opts.items()) if str(v).strip()]
        else:
            opts = [str(o).strip() for o in (raw_opts or []) if str(o).strip()]
        ans = str(q.get("answer", "")).strip().upper()
        stem = str(q.get("stem") or q.get("question") or "").strip()
        topic = str(q.get("topic", "")).strip()
        letters = [chr(65 + j) for j in range(len(opts))]
        if len(opts) < 3 or ans not in letters or not stem or topic not in QUIZ_TOPICS:
            continue
        new_questions.append({"id": next_id, "topic": topic, "stem": stem,
                              "options": dict(zip(letters, opts)), "answer": ans,
                              "explanation": str(q.get("explanation", "")).strip()})
        next_id += 1
    return jsonify({"filename": f.filename, "chars": len(text),
                    "text_preview": text[:1200],
                    "module_updates": mod_updates, "new_questions": new_questions})


if __name__ == "__main__":
    load_kb()
    load_onboarding()
    load_quiz()
    port = int(os.environ.get("PORT", "8000"))
    app.run(host="0.0.0.0", port=port, threaded=True)
